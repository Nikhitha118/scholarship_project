from flask import Blueprint, render_template, request, redirect, url_for, flash, current_app
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from app import db
from app.models.document import Document
from app.services.document_processor import DocumentProcessor

import os
from datetime import datetime
from docx import Document as DocxDocument
import re

# ✅ FIX: Blueprint name was incorrect
documents_bp = Blueprint('documents', __name__)

processor = DocumentProcessor()


# =========================
# LIST DOCUMENTS
# =========================
@documents_bp.route('/documents')
@login_required
def index():
    documents = Document.query.filter_by(user_id=current_user.id).all()
    return render_template('documents/index.html', documents=documents)


# =========================
# UPLOAD DOCUMENT
# =========================
@documents_bp.route('/documents/upload', methods=['GET', 'POST'])
@login_required
def upload():

    if request.method == 'POST':

        if 'file' not in request.files:
            flash('No file selected', 'error')
            return redirect(request.url)

        file = request.files['file']

        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)

        if file and processor.allowed_file(file.filename):

            filename = secure_filename(file.filename)

            # Unique filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_filename = f"{timestamp}_{filename}"

            # Save path
            upload_folder = os.path.join(current_app.config['UPLOAD_FOLDER'], 'profiles')
            os.makedirs(upload_folder, exist_ok=True)

            filepath = os.path.join(upload_folder, unique_filename)
            file.save(filepath)

            # =========================
            # DOCX PARSING FIXED
            # =========================
            student = {}

            if filepath.endswith('.docx'):
                try:
                    doc = DocxDocument(filepath)

                    text = "\n".join([para.text for para in doc.paragraphs])

                    print("===== DOCUMENT TEXT =====")
                    print(text)

                    # safer regex (case-insensitive + flexible spaces)
                    name_match = re.search(r'Name\s*:\s*(.*)', text, re.IGNORECASE)
                    income_match = re.search(r'Income\s*:\s*(\d+)', text, re.IGNORECASE)
                    category_match = re.search(r'Category\s*:\s*(.*)', text, re.IGNORECASE)

                    if name_match:
                        student['name'] = name_match.group(1).strip()

                    if income_match:
                        student['income'] = income_match.group(1).strip()

                    if category_match:
                        student['category'] = category_match.group(1).strip()

                    print("Parsed student:", student)
                    # Update user profile from document

                    if 'income' in student:
                        current_user.income = int(student['income'])

                    if 'category' in student:
                        current_user.category = student['category']

                    db.session.commit()

                    print("User updated:")
                    print("Income =", current_user.income)
                    print("Category =", current_user.category)

                except Exception as e:
                    print("Error reading DOCX:", str(e))
                    flash("Error processing document", "error")

            # Document type
            doc_type = request.form.get('document_type', 'other')

            # Save DB record
            document = Document(
                user_id=current_user.id,
                filename=unique_filename,
                file_path=filepath,
                document_type=doc_type,
                file_size=os.path.getsize(filepath)
            )

            db.session.add(document)
            db.session.commit()

            flash('Document uploaded successfully!', 'success')
            return redirect(url_for('documents.index'))

        else:
            flash('File type not allowed', 'error')

    return render_template('documents/upload.html')


# =========================
# VIEW DOCUMENT
# =========================
@documents_bp.route('/documents/<int:id>')
@login_required
def view(id):
    document = Document.query.get_or_404(id)

    if document.user_id != current_user.id:
        flash('Unauthorized access', 'error')
        return redirect(url_for('documents.index'))

    return render_template('documents/view.html', document=document)


# =========================
# DELETE DOCUMENT
# =========================
@documents_bp.route('/documents/<int:id>/delete', methods=['POST'])
@login_required
def delete(id):
    document = Document.query.get_or_404(id)

    if document.user_id != current_user.id:
        flash('Unauthorized access', 'error')
        return redirect(url_for('documents.index'))

    # delete file safely
    if document.file_path and os.path.exists(document.file_path):
        os.remove(document.file_path)

    db.session.delete(document)
    db.session.commit()

    flash('Document deleted successfully', 'success')
    return redirect(url_for('documents.index'))